# -*- coding: utf-8 -*-
"""Generate styled Word commercial proposal for GRC → US."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "GRC → US — коммерческое предложение.docx"

SLATE = RGBColor(0x0F, 0x17, 0x2A)
SLATE_MID = RGBColor(0x33, 0x41, 0x55)
AMBER = RGBColor(0xB4, 0x53, 0x09)
AMBER_LIGHT = RGBColor(0xFF, 0xFB, 0xEB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER = "E2E8F0"


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_table(table, header_fill: str = "0F172A", header_text: str = "FFFFFF") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                if not p.runs and p.text:
                    run = p.add_run(p.text)
                    p.text = ""
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(
                            int(header_text[0:2], 16),
                            int(header_text[2:4], 16),
                            int(header_text[4:6], 16),
                        )
            else:
                set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F8FAFC")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = SLATE if level <= 2 else AMBER if level == 3 else SLATE_MID


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11, space_after: int = 8) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = SLATE_MID
    p.paragraph_format.space_after = Pt(space_after)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = SLATE_MID


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            table.rows[r_i + 1].cells[c_i].text = val
    style_table(table)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


def add_stage_block(
    doc: Document,
    title: str,
    weeks: str,
    focus: str,
    works_sections: list[tuple[str, list[str]]],
    result: list[str],
    excludes: list[str] | None = None,
    start_note: str | None = None,
    extra_note: str | None = None,
) -> None:
    add_heading(doc, title, level=3)
    meta = doc.add_paragraph()
    r1 = meta.add_run(f"Срок: {weeks}")
    r1.bold = True
    r1.font.color.rgb = AMBER
    r1.font.size = Pt(11)
    meta.paragraph_format.space_after = Pt(6)

    add_para(doc, focus, bold=True, size=10, space_after=6)
    if start_note:
        p = doc.add_paragraph()
        r = p.add_run("Когда стартуем: ")
        r.bold = True
        r.font.color.rgb = AMBER
        r2 = p.add_run(start_note)
        r2.font.color.rgb = SLATE_MID
        r2.font.size = Pt(10)
    if extra_note:
        add_para(doc, extra_note, size=10, space_after=6)

    for section_title, items in works_sections:
        add_para(doc, section_title, bold=True, size=10, space_after=4)
        add_bullets(doc, items)

    add_para(doc, "Результат этапа", bold=True, size=10, space_after=4)
    add_bullets(doc, result)

    if excludes:
        add_para(doc, "Не в этом этапе / не обещаем", bold=True, size=10, space_after=4)
        add_bullets(doc, excludes)

    doc.add_paragraph()


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Cover band
    cover = doc.add_table(rows=1, cols=1)
    cell = cover.rows[0].cells[0]
    set_cell_shading(cell, "0F172A")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t1 = p.add_run("GRC → US\n")
    t1.font.size = Pt(28)
    t1.font.bold = True
    t1.font.color.rgb = WHITE
    t1.font.name = "Calibri"
    t2 = p.add_run("Цифровая операционная платформа\n")
    t2.font.size = Pt(14)
    t2.font.color.rgb = RGBColor(0xFB, 0xBF, 0x24)
    t2.font.name = "Calibri"
    t3 = p.add_run("Коммерческое предложение · 2026")
    t3.font.size = Pt(11)
    t3.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    t3.font.name = "Calibri"
    doc.add_paragraph()

    add_para(doc, "Исполнитель: Andrei Zobnin · архитектор системы", size=11)
    add_para(doc, "Для: GRC · выход на рынок США (выездной ремонт тяжёлого оборудования)", size=11)
    add_para(
        doc,
        "Формат: поэтапное внедрение · согласование и приёмка по этапам · сервисы — на аккаунтах GRC.",
        size=11,
    )
    add_para(
        doc,
        "Стоимость работ и схема оплаты — в приложении в конце документа, после описания этапов.",
        size=10,
        space_after=12,
    )
    doc.add_paragraph()

    # 1
    add_heading(doc, "1. Суть предложения", level=2)
    add_para(
        doc,
        "Проект — отдельный US-контур: не перестройка всего GRC в России, а своя цепочка "
        "от первого касания до повторного заказа на американском рынке.",
    )
    add_para(doc, "Полная цепочка:", bold=True, size=10, space_after=4)
    chain = doc.add_paragraph()
    cr = chain.add_run(
        "Сайт → Приём заявок → AI Estimator → CRM → (аварийный отклик) → "
        "Личный кабинет клиента → Документы → Выезд → Отчёт → Повторный заказ"
    )
    cr.bold = True
    cr.font.color.rgb = SLATE
    cr.font.size = Pt(10.5)
    add_para(
        doc,
        "AI Estimator у GRC уже есть — в проект входит подключение к US-потоку, не разработка с нуля.",
    )
    add_table(
        doc,
        ["Для встречи", "Детализация"],
        [
            ["~6 месяцев — первый рабочий контур", "24–34 недели по этапам"],
            ["Полный рабочий контур — до 8 месяцев", "с учётом параллельных работ и приёмки"],
        ],
        [2.2, 3.3],
    )
    doc.add_paragraph()

    # 2
    add_heading(doc, "2. Состав платформы", level=2)
    add_heading(doc, "Как сказать клиенту", level=3)
    quote = doc.add_paragraph()
    qr = quote.add_run(
        "В базовую версию US-платформы входит цифровой контур: заявки через сайт, "
        "email- и SMS-уведомления, личный кабинет клиента со статусами и документами. "
        "Голосовой AI-ассистент на телефоне — отдельный опциональный модуль: "
        "подключается по необходимости и не входит в базовую программу без отдельного согласования."
    )
    qr.font.size = Pt(10)
    qr.font.italic = True
    qr.font.color.rgb = SLATE_MID

    add_heading(doc, "2.1. Базовая версия платформы", level=3)
    add_table(
        doc,
        ["Компонент", "Содержание"],
        [
            ["Заявки через сайт", "Формы, аварийная заявка → CRM"],
            ["Email-уведомления", "Клиенту и команде по статусам"],
            ["SMS-уведомления", "Авария: диспетчер, автоответ, эскалация"],
            ["Личный кабинет", "Вход клиента, обзор заявок (MVP)"],
            ["Статусы", "Этап заявки для клиента и команды"],
            ["Документы", "Upload, PDF КП/отчётов; без суммы в UI"],
        ],
        [1.6, 3.6],
    )
    add_para(
        doc,
        "По программе внедрения также: AI Estimator (интеграция), CRM, контент, outreach, "
        "документы на площадку, радар, отчёты. Кабинет — этап 3 после стабильного intake с сайта.",
        size=9,
    )

    add_heading(doc, "2.2. Опциональный модуль (не в базе)", level=3)
    add_table(
        doc,
        ["Модуль", "Описание"],
        [
            [
                "Голосовой AI-ассистент",
                "Телефон 24/7: сбор заявки, planned/emergency → CRM. Не входит в $43 000 без отдельного согласования.",
            ],
        ],
        [1.4, 4.8],
    )

    add_heading(doc, "2.3. Не входит в программу", level=3)
    add_bullets(
        doc,
        [
            "разработка AI Estimator с нуля",
            "замена CRM или бухгалтерии GRC",
            "гарантия фиксированного числа сделок",
            "перестройка 1grc.ru",
        ],
    )
    doc.add_paragraph()

    # 3
    add_heading(doc, "3. Этапы внедрения", level=2)
    add_para(
        doc,
        "Четыре этапа с понятной сдачей. Сроки — в неделях. Календарь зависит от согласований и доступов GRC.",
        size=10,
    )
    add_para(
        doc,
        "Ориентир по программе: ~6 месяцев до рабочего контура; полный контур — до 8 месяцев.",
        size=10,
        space_after=10,
    )

    add_stage_block(
        doc,
        "Этап 1 · Фундамент: заявка не теряется",
        "6–10 недель",
        "Фокус: модули 1–3 · приём заявок, аварийный отклик, CRM",
        [
            ("Недели 1–2 — подготовка", [
                "воронка CRM под US B2B-цикл",
                "роли: диспетчер, ночной ответственный",
                "шаблоны SMS и автоответов",
                "формат карточки для инженера и AI Estimator",
            ]),
            ("Недели 3–6 — первая сдача", [
                "US-сайт: заявка и кнопка аварии → CRM",
                "авария: SMS, автоответ, эскалация",
                "связка заявка → AI Estimator → CRM",
                "регламент ночного отклика (15 мин)",
            ]),
        ],
        [
            "заявка с сайта в CRM, не в почте",
            "на аварию — SMS и ответ клиенту",
            "Estimator получает структурированный пакет",
        ],
        [
            "голосовой AI-ассистент (опциональный модуль)",
            "личный кабинет — этап 3",
            "радар остановок и полный пакет документов",
            "контент и массовый прямой выход",
        ],
    )

    add_stage_block(
        doc,
        "Этап 2 · Вас находят и выходите на людей",
        "6–8 недель",
        "Фокус: модули 4–5 · контент и прямой выход",
        [
            ("Работы", [
                "15–20 кейсов с 1grc.ru → US-сайт",
                "базовое SEO по поломке / ремонту на объекте",
                "база контактов Gulf Coast",
                "шаблоны писем под оборудование на площадке",
                "прямой выход: 50–100 контактов, трекинг в CRM",
                "ежемесячный отчёт по воронке",
            ]),
        ],
        [
            "US-инженер находит GRC по задаче",
            "выход на тех, кто подписывает Purchase order",
        ],
        ["фиксированное число сделок — обещаем процесс в CRM"],
        start_note="параллельно концу этапа 1 — после первой сдачи сайта и CRM",
    )

    add_stage_block(
        doc,
        "Этап 3 · Личный кабинет клиента (MVP)",
        "4–6 недель",
        "Фокус: личный кабинет · слой поверх CRM",
        [
            ("Работы", [
                "один контакт на компанию",
                "дашборд: заявки, статусы, что дальше",
                "карточка: объём работ + шаг (без суммы в UI)",
                "документы: upload, PDF КП и отчётов",
                "повторная заявка → CRM",
            ]),
        ],
        [
            "клиент видит статус без переписки",
            "воспринимается как платформа, не «форма»",
        ],
        ["чат", "цены в UI", "несколько пользователей на компанию"],
        start_note="после 2+ недель стабильного приёма заявок в CRM",
    )

    add_stage_block(
        doc,
        "Этап 4 · Крупные контракты и повторные заказы",
        "8–10 недель",
        "Фокус: модули 6–8 · документы, радар, отчёты, память",
        [
            ("Работы", [
                "пакет COI, OSHA, safety — сборка и напоминания",
                "радар turnarounds: сигнал за 3–6 месяцев",
                "отчёт PDF за 48 часов",
                "память объекта в CRM/базе",
                "внутренняя панель над CRM",
                "5–7 сценариев автоматизации (согласование на старте)",
            ]),
        ],
        [
            "контур: касание → выезд → документы → отчёт → повтор",
            "опыт не уезжает с сотрудником",
        ],
        ["замена CRM", "enterprise-аналитика вне списка workflow"],
    )

    add_heading(doc, "4. Приёмка этапа", level=2)
    add_para(doc, "По согласованному чеклисту, не по календарной дате.", size=10)
    add_bullets(
        doc,
        [
            "Этап 1: заявка в CRM, авария → SMS, сквозной кейс Estimator, регламент GRC",
            "Этап 2: кейсы live, воронка в CRM, отчёт за месяц",
            "Этап 3: тестовый клиент до повторной заявки",
            "Этап 4: чеклист 5–7 пунктов на kickoff",
        ],
    )

    add_heading(doc, "5. Старт этапа 1 — что нужно от GRC", level=2)
    add_bullets(
        doc,
        [
            "Подтверждение отдельного US pipeline",
            "Диспетчер на аварию ночью (SMS, эскалация)",
            "AI Estimator: что работает, формат данных",
            "Выбор CRM",
            "US-сайт: приоритет production",
            "Single point of contact у клиента (для кабинета)",
        ],
    )

    # 6 — расходы GRC (их инфраструктура, не гонорар исполнителя)
    add_heading(doc, "6. Расходы GRC на инфраструктуру", level=2)
    add_para(
        doc,
        "Оплачиваются напрямую GRC со своих аккаунтов — отдельно от работ исполнителя.",
        size=10,
    )
    add_heading(doc, "6.1. Ежемесячно (ориентир на старте)", level=3)
    add_table(
        doc,
        ["Категория", "USD/мес"],
        [
            ["Хостинг и публикация", "~30"],
            ["База данных, авторизация, файлы", "~50"],
            ["Платформа автоматизаций", "~70"],
            ["AI API", "~200"],
            ["Email / SMS", "~100"],
            ["CRM", "~150"],
            ["Домены и почта", "~50"],
            ["Итого ориентировочно", "~650"],
        ],
        [4.0, 1.2],
    )
    add_para(doc, "При рабочей нагрузке: ориентир ~1 500–2 500 / мес.", size=10)
    add_heading(doc, "6.2. Разовый старт", level=3)
    add_table(
        doc,
        ["Позиция", "USD"],
        [
            ["Домены", "~100–200"],
            ["Стартовые кредиты AI API", "~300–500"],
            ["Настройка email/SMS", "~100–300"],
            ["Итого ориентировочно", "~500–1 000"],
        ],
        [4.0, 1.2],
    )

    # Приложение — стоимость (скромно, в конце)
    doc.add_page_break()
    add_heading(doc, "Приложение · Стоимость работ и оплата", level=2)
    add_para(
        doc,
        "Ниже — ориентиры по бюджету работ исполнителя. Детализация по этапам; оплата после приёмки "
        "каждого этапа по чеклисту. Можно начать только с этапа 1.",
        size=10,
        space_after=10,
    )

    add_heading(doc, "Стоимость по этапам", level=3)
    add_table(
        doc,
        ["Этап", "Срок", "Стоимость работ"],
        [
            ["1 · Фундамент", "6–10 нед", "$13 000"],
            ["2 · Видимость и выход", "6–8 нед", "$10 000"],
            ["3 · Личный кабинет", "4–6 нед", "$9 000"],
            ["4 · Масштаб и повтор", "8–10 нед", "$11 000"],
        ],
        [2.8, 1.1, 1.2],
    )
    p_total = doc.add_paragraph()
    p_total.paragraph_format.space_before = Pt(8)
    r = p_total.add_run("Итого по программе (4 этапа): ")
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_MID
    r2 = p_total.add_run("$43 000")
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = SLATE
    add_para(
        doc,
        "Срок программы: ~7–9 месяцев по календарю; на встрече — ориентир ~6 месяцев, "
        "полный контур до ~8 месяцев. Входной вариант: только этап 1 — $13 000.",
        size=9,
        space_after=12,
    )

    add_heading(doc, "Схема оплаты", level=3)
    add_para(doc, "По каждому этапу: 50% на старт, 50% после приёмки.", size=10, space_after=6)
    add_table(
        doc,
        ["Этап", "Старт (50%)", "Приёмка (50%)", "Итого"],
        [
            ["1", "$6 500", "$6 500", "$13 000"],
            ["2", "$5 000", "$5 000", "$10 000"],
            ["3", "$4 500", "$4 500", "$9 000"],
            ["4", "$5 500", "$5 500", "$11 000"],
            ["Всего", "$21 500", "$21 500", "$43 000"],
        ],
        [0.55, 1.15, 1.15, 1.0],
    )

    add_heading(doc, "Поддержка после запуска (опционально)", level=3)
    add_para(
        doc,
        "Базовый пакет — ориентир ~$1 500/мес: мониторинг, мелкие правки сценариев, "
        "настройка AI, до 5 часов доработок, ежемесячный обзор.",
        size=10,
    )
    add_para(
        doc,
        "Расширенный — ориентир ~$3 000/мес: до 20 часов улучшений, приоритетная поддержка. "
        "Отдельное соглашение; крупные новые модули — отдельный этап.",
        size=10,
        space_after=12,
    )

    add_heading(doc, "Сводка бюджетов", level=3)
    add_table(
        doc,
        ["Статья", "Ориентир"],
        [
            ["Работы исполнителя (программа)", "$43 000"],
            ["Инфраструктура GRC, разово", "~$500–1 000"],
            ["Инфраструктура GRC, ежемесячно (старт)", "~$650 / мес"],
            ["Инфраструктура GRC, при нагрузке", "~$1 500–2 500 / мес"],
            ["Поддержка (если нужна)", "~$1 500 или ~$3 000 / мес"],
        ],
        [3.2, 2.0],
    )

    doc.add_paragraph()
    add_heading(doc, "Контакты исполнителя", level=2)
    add_table(
        doc,
        ["", ""],
        [
            ["Имя", "Andrei Zobnin"],
            ["Роль", "Архитектор системы · GRC → US"],
            ["Email", "zobnin8@gmail.com"],
            ["Telegram", "@A_nubi_ss"],
            ["Сайт", "https://www.zobnin.tech/"],
        ],
        [1.5, 4.0],
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Документ 2026-05-26 · Согласован с pitch-презентацией GRC → US")
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUT)
    print("Saved:", OUT.name)


if __name__ == "__main__":
    build()
