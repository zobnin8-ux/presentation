# -*- coding: utf-8 -*-
"""Generate Word doc: GRC US technical architecture and subscriptions."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "GRC → US — техархитектура и подписки.docx"

SLATE = RGBColor(0x0F, 0x17, 0x2A)
SLATE_MID = RGBColor(0x33, 0x41, 0x55)
AMBER = RGBColor(0xB4, 0x53, 0x09)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_table(table, header_fill: str = "0F172A") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                if not p.runs and p.text:
                    run = p.add_run(p.text)
                    p.text = ""
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
            else:
                set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F8FAFC")


def heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = SLATE if level <= 2 else AMBER


def para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = SLATE_MID
    p.paragraph_format.space_after = Pt(8)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    style_table(t)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)


def build() -> None:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

    cov = doc.add_table(rows=1, cols=1)
    cell = cov.rows[0].cells[0]
    set_cell_shading(cell, "0F172A")
    cell.text = ""
    p = cell.paragraphs[0]
    t1 = p.add_run("GRC → US\n")
    t1.font.size = Pt(24)
    t1.font.bold = True
    t1.font.color.rgb = WHITE
    t1.font.name = "Calibri"
    t2 = p.add_run("Техническая архитектура и подписки\n")
    t2.font.size = Pt(13)
    t2.font.color.rgb = RGBColor(0xFB, 0xBF, 0x24)
    t2.font.name = "Calibri"
    t3 = p.add_run("Внутренний документ · Andrei Zobnin · 2026")
    t3.font.size = Pt(10)
    t3.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    t3.font.name = "Calibri"
    doc.add_paragraph()

    para(
        doc,
        "Только для исполнителя. Не для GRC. Подписки GRC — на их аккаунтах. "
        "Раздел 4 — ваши фиксированные расходы на инструменты.",
        size=10,
    )

    heading(doc, "1. Что строится", 2)
    para(doc, "Отдельный US pipeline (не перестройка 1grc.ru):", bold=True, size=10)
    flow = doc.add_paragraph()
    fr = flow.add_run(
        "US-сайт + телефон → приём 24/7 → AI Estimator (у GRC) → CRM → "
        "документы → выезд → отчёт → личный кабинет → повтор. "
        "Склейка: n8n/Make + Postgres (Supabase/Neon)."
    )
    fr.font.size = Pt(10)
    fr.font.color.rgb = SLATE_MID
    para(doc, "Репозитории: presentation (pitch) · grc / US-сайт.", size=10)

    heading(doc, "2. Архитектура по слоям", 2)
    table(
        doc,
        ["Слой", "Технология", "Этап"],
        [
            ["Витрина", "Next.js 15, Vercel", "1"],
            ["Intake сайт", "API, Resend", "1"],
            ["Intake телефон", "Twilio Voice", "1 (волна 2)"],
            ["Emergency", "Twilio SMS, n8n", "1"],
            ["AI Estimator", "Существующий у GRC", "1"],
            ["CRM", "Pipedrive / HubSpot", "1"],
            ["Автоматизация", "n8n", "1–4"],
            ["БД / auth / files", "Supabase / Neon", "1–3"],
            ["Личный кабинет", "Auth + dashboard + PDF", "3"],
            ["Документы, отчёты", "Drive + шаблоны, PDF", "4"],
            ["Outreach, радар", "CRM, LinkedIn, Apollo", "2–4"],
        ],
        [1.4, 2.2, 0.9],
    )

    heading(doc, "3. Production — платит GRC", 2)
    table(
        doc,
        ["Сервис", "USD/мес (ориентир)"],
        [
            ["Twilio (SMS, voice)", "30–150"],
            ["CRM (1–3 места)", "45–270"],
            ["Google Workspace", "12–42"],
            ["AI API (prod)", "100–300"],
            ["Supabase/Neon, n8n, email, домен", "20–120"],
            ["Итого на старте", "~200–650 (в КП ~650)"],
            ["При нагрузке", "~1 500–2 500"],
        ],
        [3.5, 1.8],
    )
    para(doc, "Разовый старт GRC: ~$500–1 000 (домены, кредиты AI, SMS).", size=10)

    heading(doc, "4. Подписки исполнителя (ваши)", 2)
    para(
        doc,
        "Не выставляются GRC — входят в гонорар $43 000. Бюджет на весь проект ниже.",
        size=10,
    )
    table(
        doc,
        ["Сервис", "Назначение", "USD/мес"],
        [
            ["Cursor", "IDE + AI", "60"],
            ["OpenAI", "API, баланс ≥", "50"],
            ["Anthropic", "API, баланс ≥", "50"],
            ["Claude (Chat)", "Подписка claude.ai, отдельно от API", "100"],
            ["Vercel", "Preview, деплои, serverless", "100–200"],
            ["Итого", "", "360–460"],
        ],
        [1.2, 2.4, 1.0],
    )
    para(doc, "Рекомендуем планировать: ~$400/мес (Vercel ~$150, Claude Chat $100).", bold=True, size=10)

    heading(doc, "4.1. За весь проект (7–9 мес)", 3)
    table(
        doc,
        ["Сценарий", "USD/мес", "Мес.", "Итого"],
        [
            ["Минимум (Vercel $100)", "360", "8", "~2 880"],
            ["Базовый (Vercel $150)", "410", "8", "~3 280"],
            ["Максимум (Vercel $200)", "460", "9", "~4 140"],
        ],
        [1.8, 0.7, 0.5, 1.0],
    )
    para(doc, "Ориентир: ~$3 000–4 000 на подписки за полный цикл.", bold=True, size=10)

    heading(doc, "5. Кто платит", 2)
    table(
        doc,
        ["", "GRC", "Вы"],
        [
            ["CRM, Twilio prod, домен", "✓", ""],
            ["AI API prod", "✓", ""],
            ["Cursor, OpenAI/Anthropic API, Claude Chat", "", "✓"],
            ["Vercel dev/preview", "", "✓"],
            ["Работы по этапам", "", "$43 000"],
        ],
        [2.5, 0.6, 0.6],
    )

    heading(doc, "6. Полная картина", 2)
    table(
        doc,
        ["Статья", "Сумма"],
        [
            ["Гонорар (КП)", "$43 000"],
            ["Ваши подписки (~8 мес)", "~$3 000–4 000"],
            ["GRC подписки старт", "~$650/мес"],
            ["GRC подписки нагрузка", "~$1 500–2 500/мес"],
        ],
        [3.0, 2.2],
    )

    heading(doc, "7. TL;DR", 2)
    para(
        doc,
        "Вы каждый месяц: Cursor $60 + OpenAI API ≥$50 + Anthropic API ≥$50 + Claude Chat $100 + "
        "Vercel $100–200 = $360–460. За проект на подписки: ~$3 000–4 000. GRC платит prod отдельно. "
        "Ваш гонорар $43 000 — отдельно от их Twilio/CRM.",
        size=10,
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Конфиденциально · не для клиента · 2026-05-26")
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print("OK")
